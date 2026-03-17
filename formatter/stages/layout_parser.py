"""
stages/layout_parser.py
Stage 1 — PDF/DOCX → raw text blocks (extracted.txt)

PDF path:  pix2text  → layout + formula detection + OCR → Markdown
           Falls back to PyMuPDF if pix2text is not installed.

DOCX path: python-docx → paragraphs, equations, tables (unchanged)
"""

import os
import re


def parse(input_path: str, output_path: str) -> str:
    ext = os.path.splitext(input_path)[1].lower()

    if ext == ".pdf":
        text = _from_pdf(input_path)
    elif ext in (".docx", ".doc"):
        text = _from_docx(input_path)
    else:
        raise ValueError(f"Unsupported file type '{ext}'. Expected .pdf or .docx")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"         → {output_path}")
    return output_path


# ── PDF: pix2text (preferred) with PyMuPDF fallback ──────────────────────────

def _from_pdf(path: str) -> str:
    try:
        return _from_pdf_pix2text(path)
    except ImportError:
        print("         [layout_parser] pix2text not found — falling back to PyMuPDF")
        print("         Install with: pip install pix2text")
        return _from_pdf_pymupdf(path)


def _from_pdf_pix2text(path: str) -> str:
    """
    Use pix2text to extract PDF content with formula recognition.
    Returns double-newline-separated blocks for document_parser.
    """
    from pix2text import Pix2Text

    print("         [pix2text] loading models (first run downloads ~500MB)...")
    p2t = Pix2Text.from_config()

    print(f"         [pix2text] recognising {os.path.basename(path)}...")
    pages = p2t.recognize_pdf(path, page_numbers=None)

    blocks = []
    for page in pages:
        md = page.to_markdown(None)
        if md:
            for para in _md_to_blocks(md):
                if para.strip():
                    blocks.append(para.strip())

    return "\n\n".join(blocks)


def _md_to_blocks(md: str) -> list[str]:
    """
    Split pix2text Markdown into double-newline-separated blocks,
    keeping $$...$$ display math blocks intact.
    """
    blocks          = []
    current         = []
    in_display_math = False

    for line in md.split("\n"):
        stripped = line.strip()

        if stripped == "$$":
            if in_display_math:
                current.append("$$")
                blocks.append("\n".join(current))
                current = []
                in_display_math = False
            else:
                if current:
                    blocks.append("\n".join(current))
                    current = []
                current.append("$$")
                in_display_math = True
            continue

        if in_display_math:
            current.append(line)
            continue

        if stripped == "":
            if current:
                blocks.append("\n".join(current))
                current = []
        else:
            current.append(line)

    if current:
        blocks.append("\n".join(current))

    return blocks


def _from_pdf_pymupdf(path: str) -> str:
    """Fallback: PyMuPDF plain text + image extraction."""
    import fitz
    from core.config import INTERMEDIATE_DIR

    figures_dir = os.path.join(INTERMEDIATE_DIR, "figures")
    os.makedirs(figures_dir, exist_ok=True)

    doc        = fitz.open(path)
    paragraphs = []
    fig_count  = 0

    for page_num, page in enumerate(doc):
        # ── Extract images on this page ───────────────────────────────────────
        image_list = page.get_images(full=True)
        page_images = []  # (y_position, filepath) for sorting with text

        for img_index, img in enumerate(image_list):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                img_bytes  = base_image["image"]
                img_ext    = base_image["ext"]

                # Skip tiny images (icons, bullets, decorations) — under 5KB
                if len(img_bytes) < 5000:
                    continue

                fig_count += 1
                img_filename = f"fig_{fig_count}.{img_ext}"
                img_path     = os.path.join(figures_dir, img_filename)

                with open(img_path, "wb") as f:
                    f.write(img_bytes)

                # Get approximate y-position of image on page for ordering
                img_rects = page.get_image_rects(xref)
                y_pos = img_rects[0].y0 if img_rects else 0

                # Emit a special marker so document_parser can place a figure
                rel_path = os.path.join("intermediate", "figures", img_filename).replace("\\", "/")
                marker = f"__IMAGE__{rel_path}__PAGE_{page_num}__Y_{int(y_pos)}__"
                page_images.append((y_pos, marker))

            except Exception:
                continue

        # ── Extract text blocks on this page ─────────────────────────────────
        blocks = page.get_text("blocks")
        blocks.sort(key=lambda b: (round(b[1] / 20) * 20, b[0]))

        # Merge text blocks and image markers in vertical order
        text_items = []
        for block in blocks:
            text = block[4].strip()
            if text:
                text_items.append((block[1], text))   # (y0, text)

        for y_pos, marker in page_images:
            text_items.append((y_pos, marker))

        text_items.sort(key=lambda x: x[0])
        paragraphs.extend(item[1] for item in text_items)

    return "\n\n".join(paragraphs)


# ── DOCX extraction (unchanged) ───────────────────────────────────────────────

def _from_docx(path: str) -> str:
    from docx import Document as DocxDocument

    doc        = DocxDocument(path)
    body       = doc.element.body
    para_idx   = 0
    table_idx  = 0
    output_blocks: list[str] = []

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if para_idx < len(doc.paragraphs):
                para     = doc.paragraphs[para_idx]
                para_idx += 1
                text     = para.text.strip()
                if not text:
                    continue

                # Equation: [\n\\latex\n]
                eq_match = re.match(r"^\[\s*\n?(.*?)\n?\s*\]$", text, re.DOTALL)
                if eq_match:
                    latex = eq_match.group(1).strip().replace("\\\\", "\\")
                    output_blocks.append(f"$$\n{latex}\n$$")
                    continue

                # Multi-author block near top of document
                if "\n" in text and para_idx <= 5:
                    output_blocks.append(text)
                    continue

                output_blocks.append(text)

        elif tag == "tbl":
            if table_idx < len(doc.tables):
                tbl       = doc.tables[table_idx]
                table_idx += 1

                caption = ""
                for prev in reversed(output_blocks[-3:]):
                    if len(prev.split()) <= 6 and not prev.startswith("$$"):
                        caption = prev
                        output_blocks.remove(prev)
                        break

                rows = []
                for row in tbl.rows:
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    rows.append(cells)

                if rows:
                    n_cols    = max(len(r) for r in rows)
                    header    = (rows[0] + [""] * n_cols)[:n_cols]
                    tbl_lines = [" & ".join(header)]
                    for row in rows[1:]:
                        tbl_lines.append(" & ".join((row + [""] * n_cols)[:n_cols]))
                    cap_text = caption if caption else f"Table {table_idx}"
                    output_blocks.append(
                        "DOCX_TABLE_START\n" +
                        "\n".join(tbl_lines) +
                        f"\nTable {table_idx}: {cap_text}\nDOCX_TABLE_END"
                    )

    return "\n\n".join(output_blocks)