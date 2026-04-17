"""
extractor/docx_extractor.py — DOCX text extraction via python-docx.
"""
import os
import logging

log = logging.getLogger("paper_formatter")


def extract_docx(path: str) -> dict:
    """Extract text and tables from a DOCX file."""
    from docx import Document as DocxDocument

    log.info("  Opening DOCX: %s", os.path.basename(path))
    doc = DocxDocument(path)

    # Extract paragraphs
    paragraphs = []
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)

        # Get font info from first run
        font_name = ""
        font_size = 0
        if para.runs:
            run = para.runs[0]
            if run.font.name:
                font_name = run.font.name
            if run.font.size:
                font_size = round(run.font.size.pt, 1)

        # Detect heading style
        is_heading = para.style and para.style.name.startswith("Heading")

        blocks.append({
            "text": text,
            "font": font_name,
            "size": font_size,
            "page": 0,
            "bbox": [0, 0, 0, 0],
            "is_heading": is_heading,
            "heading_level": int(para.style.name[-1]) if is_heading and para.style and para.style.name and para.style.name[-1].isdigit() else 0,
        })

    # Extract tables
    tables = []
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if len(rows_data) >= 2:
            tables.append({
                "headers": rows_data[0],
                "rows": rows_data[1:],
                "caption": "",
                "label": f"tab_docx_{i+1}",
            })

    full_text = "\n".join(paragraphs)
    log.info("  Extracted %d chars, %d blocks, %d tables from DOCX",
             len(full_text), len(blocks), len(tables))

    return {
        "text": full_text,
        "blocks": blocks,
        "tables": tables,
        "figures": [],
        "formula_blocks": [],
    }
